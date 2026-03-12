# app/document_processor.py
"""
Simple document processor for PDFs and DOCX files.
"""

import os
from pathlib import Path
from PyPDF2 import PdfReader

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    if DocxDocument is None:
        return ""
    
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file (text-based)."""
    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""


def extract_text(file_path: str) -> str:
    """
    Extract text from file based on extension.
    Supports: .pdf, .docx
    """
    ext = Path(file_path).suffix.lower()
    
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        return ""