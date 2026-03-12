# app/utils.py
import os
import base64
import time
import tempfile
from typing import List
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from openai import OpenAI

from app.config import settings

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None


# ============================================================
# 📄 PDF Text Extraction
# ============================================================

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts text from a PDF file.
    First tries PyPDF2 (fast for text PDFs).
    If that fails, uses gpt-4o with vision for scanned PDFs.
    """
    text = ""
    
    # Try PyPDF2 first (fast)
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
        
        if text.strip():
            return text.strip()
    except:
        pass
    
    # If PyPDF2 got nothing, use gpt-4o for scanned PDFs
    print(f"  → Scanned PDF detected, using gpt-4o...", end=" ")
    return extract_scanned_pdf_gpt4o(file_path)


def extract_scanned_pdf_gpt4o(file_path: str) -> str:
    """
    Extract text from scanned PDF using gpt-4o vision.
    Converts PDF pages to images and uses gpt-4o to read them.
    Only uses your OpenAI API key (no extra Vision API needed).
    """
    if not convert_from_path:
        print("⚠️  pdf2image not installed")
        return ""
    
    try:
        # Convert PDF to images
        images = convert_from_path(file_path, dpi=150)
        
        if not images:
            return ""
        
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        all_text = []
        
        # Process each page
        for idx, image in enumerate(images[:10]):  # Limit to first 10 pages
            try:
                # Save image to temp file
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    image.save(tmp.name)
                    
                    # Encode to base64
                    with open(tmp.name, "rb") as f:
                        image_data = base64.standard_b64encode(f.read()).decode("utf-8")
                    
                    # Use gpt-4o (includes vision) to extract text
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/png;base64,{image_data}"
                                        }
                                    },
                                    {
                                        "type": "text",
                                        "text": "Extract all text from this image exactly as written. Return only the text content."
                                    }
                                ]
                            }
                        ],
                        max_tokens=1500
                    )
                    
                    page_text = response.choices[0].message.content
                    if page_text:
                        all_text.append(page_text)
                    
                    # Clean up
                    try:
                        os.remove(tmp.name)
                    except:
                        pass
                    time.sleep(0.3)
            
            except Exception as e:
                print(f"Error on page {idx+1}: {e}")
                continue
        
        return "\n".join(all_text).strip()
    
    except Exception as e:
        print(f"Error: {e}")
        return ""


# ============================================================
# 📄 DOCX Text Extraction
# ============================================================

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file.
    Returns concatenated string of paragraphs and tables.
    """
    if DocxDocument is None:
        return ""
    
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""


# ============================================================
# ✂️ Smart Chunking
# ============================================================

def smart_chunk_text(text: str) -> List[str]:
    """
    Splits text into semantically-aware overlapping chunks.
    Uses RecursiveCharacterTextSplitter from LangChain.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", "?", "!", " ", ""],
    )
    return splitter.split_text(text)


# ============================================================
# 🧠 Embedding & Vectorstore
# ============================================================

def create_faiss_index(chunks: List[str], metadata: dict = None) -> FAISS:
    """
    Create a FAISS vector store from text chunks and metadata.
    Saves it to local path (settings.VECTOR_STORE_PATH).
    """
    embeddings = OpenAIEmbeddings(
        model=settings.EMBEDDING_MODEL,
        openai_api_key=settings.OPENAI_API_KEY,
    )
    vectorstore = FAISS.from_texts(chunks, embedding=embeddings, metadatas=[metadata or {}] * len(chunks))
    save_faiss_index(vectorstore, settings.VECTOR_STORE_PATH)
    return vectorstore


def save_faiss_index(vectorstore: FAISS, path: str):
    """Saves FAISS index to disk."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    vectorstore.save_local(path)


def load_faiss_index(path: str) -> FAISS:
    """Loads FAISS index from disk. Returns None if not found."""
    if not os.path.exists(path):
        return None
    embeddings = OpenAIEmbeddings(
        model=settings.EMBEDDING_MODEL,
        openai_api_key=settings.OPENAI_API_KEY,
    )
    return FAISS.load_local(path, embeddings, allow_dangerous_deserialization=True)


# ============================================================
# 📊 Utility Helpers
# ============================================================

def count_unique_clients(metadata_list: List[dict]) -> int:
    """Counts unique client names from a list of metadata dicts."""
    names = {m.get("client_name", "").strip().lower() for m in metadata_list if m.get("client_name")}
    return len(names)


def count_unique_regions(metadata_list: List[dict]) -> int:
    """Counts unique regions across all contracts."""
    regions = {m.get("region", "").strip().lower() for m in metadata_list if m.get("region")}
    return len(regions)